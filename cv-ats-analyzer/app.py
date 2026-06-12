"""
CV ATS Analyzer — Flask Application
Zaawansowany ATS do analizy CV dla stanowisk technicznych Senior PHP Developer.
"""

import os
import re
import json
import html
import datetime
import urllib.parse
import urllib.request
from pathlib import Path

from flask import Flask, render_template, request, jsonify, send_file
from services.analysis_service import (
    build_improvement_annotations as analysis_build_improvement_annotations,
    build_text_summary as analysis_build_text_summary,
    run_ats_analysis as analysis_run_ats_analysis,
)
from services.document_service import (
    extract_pdf_layout as document_extract_pdf_layout,
    extract_text as document_extract_text,
)
from services.preview_service import (
    build_html_from_layout as preview_build_html_from_layout,
    convert_faithful_template as preview_convert_faithful_template,
    export_html_preview_file as preview_export_html_preview_file,
    inline_preview_assets as preview_inline_preview_assets,
    list_html_preview_items,
    read_html_preview_content,
)
from services.profile_service import ATS_PROFILE, POPULAR_JOB_DOMAINS, TEMPLATE_CATALOG
from services.storage import (
    ALLOWED_EXTENSIONS,
    GENERATED_DIR,
    HTML_PREVIEW_DIR,
    UPLOAD_DIR,
    ensure_directories,
)
from services.storage import (
    CONVERSION_ASSETS_DIR,
    EXTRACTED_IMAGES_DIR,
)

ensure_directories()
# ── Manifest helpers ──────────────────────────────────────────────────────────
MANIFEST_FILE = GENERATED_DIR / "preview_manifest.json"

def _manifest_load() -> dict:
    try:
        return json.loads(MANIFEST_FILE.read_text(encoding="utf-8")) if MANIFEST_FILE.exists() else {}
    except Exception:
        return {}

def _manifest_save(data: dict) -> None:
    try:
        MANIFEST_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass

def _manifest_register(upload_filename: str, preview_html: str, ts: str) -> None:
    data = _manifest_load()
    entry = data.setdefault(upload_filename, {"previews": [], "ts": []})
    if preview_html not in entry["previews"]:
        entry["previews"].append(preview_html)
    if ts not in entry["ts"]:
        entry["ts"].append(ts)
    _manifest_save(data)


app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB
def extract_text(filepath: Path) -> str:
    return document_extract_text(filepath)


def run_ats_analysis(cv_text: str, job_description: str = "") -> dict:
    return analysis_run_ats_analysis(cv_text, job_description)


def build_improvement_annotations(cv_text: str) -> list[dict]:
    return analysis_build_improvement_annotations(cv_text)


def build_text_summary(cv_text: str, analysis: dict, annotations: list[dict]) -> dict:
    return analysis_build_text_summary(cv_text, analysis, annotations)


def improve_cv_text(cv_text: str, analysis: dict, job_description: str = "") -> str:
    improved = cv_text

    rewrites = [
        (r"\bodpowiedzialn\w*\s+za\b", "Zrealizowałem"),
        (r"\buczestniczy\w*\s+w\b", "Wdrożyłem"),
        (r"\bznajomo[śs]ć\s+([A-Za-z0-9#+.\-/ ]+)", r"Praktyczne doświadczenie komercyjne: \1"),
        (r"\bteam player\b", "Skuteczna współpraca zespołowa poparta mentoringiem i code review"),
    ]

    for pattern, replacement in rewrites:
        improved = re.sub(pattern, replacement, improved, flags=re.IGNORECASE)

    must_missing = analysis["keywords"]["must_have"]["missing"][:8]
    jd_missing = analysis.get("job_context", {}).get("focus_terms", {}).get("missing", [])[:6]
    extra_terms = []
    for term in must_missing + jd_missing:
        if term not in extra_terms:
            extra_terms.append(term)

    additions = []
    if extra_terms:
        additions.append("\n\n### Dodatkowe frazy ATS (do wplecenia w doświadczenie)\n")
        additions.extend([f"- {term}" for term in extra_terms])

    recs = analysis.get("recommendations", [])[:4]
    if recs:
        additions.append("\n\n### Rekomendowane zdania (wersja poprawiona)\n")
        additions.extend([f"- {r}" for r in recs])

    if job_description.strip():
        additions.append("\n\n### Kontekst oferty\n")
        additions.append("- CV zostało zaktualizowane pod dostarczoną ofertę pracy.")

    if additions:
        improved += "\n" + "\n".join(additions)

    return improved


def improve_line_text(line_text: str, analysis: dict) -> str:
    improved = line_text
    rewrites = [
        (r"\bodpowiedzialn\w*\s+za\b", "Zrealizowałem"),
        (r"\buczestniczy\w*\s+w\b", "Wdrożyłem"),
        (r"\bznajomo[śs]ć\s+([A-Za-z0-9#+.\-/ ]+)", r"Doświadczenie komercyjne: \1"),
        (r"\bteam player\b", "Skuteczna współpraca zespołowa"),
    ]
    for pattern, replacement in rewrites:
        improved = re.sub(pattern, replacement, improved, flags=re.IGNORECASE)

    max_len = max(18, int(len(line_text) * 1.35))
    if len(improved) > max_len:
        improved = improved[: max_len - 1].rstrip() + "…"
    return improved


def extract_pdf_layout(filepath: Path) -> list[dict]:
    return document_extract_pdf_layout(filepath)


def build_html_from_layout(pages_layout: list[dict], analysis: dict, ts: str) -> tuple[str, Path]:
    return preview_build_html_from_layout(pages_layout, analysis, ts)


def render_pdf_from_layout(pages_layout: list[dict], output_path: Path):
    from reportlab.pdfgen import canvas
    from reportlab.lib import colors

    font_regular, _, _ = _build_unicode_fonts_for_pdf()

    if not pages_layout:
        c = canvas.Canvas(str(output_path))
        c.setFont(font_regular, 11)
        c.drawString(50, 780, "Brak treści do wygenerowania.")
        c.save()
        return

    first = pages_layout[0]
    c = canvas.Canvas(str(output_path), pagesize=(first["width"], first["height"]))

    for page in pages_layout:
        c.setPageSize((page["width"], page["height"]))
        c.setStrokeColor(colors.HexColor("#e1e4e8"))
        c.rect(8, 8, page["width"] - 16, page["height"] - 16, stroke=1, fill=0)

        for line in page["lines"]:
            y = page["height"] - float(line["top"]) - float(line["font_size"])
            c.setFont(font_regular, float(line["font_size"]))
            c.setFillColor(colors.black)
            c.drawString(float(line["x"]), max(12, y), str(line["text"]))

        c.showPage()
    c.save()


def _build_unicode_fonts_for_pdf():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_regular = "Helvetica"
    font_bold = "Helvetica-Bold"
    font_italic = "Helvetica-Oblique"
    font_candidates = [
        ("ATSUnicodeRegular", "/System/Library/Fonts/Supplemental/Arial.ttf"),
        ("ATSUnicodeBold", "/System/Library/Fonts/Supplemental/Arial Bold.ttf"),
        ("ATSUnicodeItalic", "/System/Library/Fonts/Supplemental/Arial Italic.ttf"),
        ("ATSUnicodeRegular", "/Library/Fonts/Arial Unicode.ttf"),
        ("ATSUnicodeRegular", "/Library/Fonts/DejaVuSans.ttf"),
        ("ATSUnicodeBold", "/Library/Fonts/DejaVuSans-Bold.ttf"),
        ("ATSUnicodeItalic", "/Library/Fonts/DejaVuSans-Oblique.ttf"),
    ]
    for name, path in font_candidates:
        if Path(path).exists() and name not in pdfmetrics.getRegisteredFontNames():
            try:
                pdfmetrics.registerFont(TTFont(name, path))
            except Exception:
                pass

    if "ATSUnicodeRegular" in pdfmetrics.getRegisteredFontNames():
        font_regular = "ATSUnicodeRegular"
        font_bold = "ATSUnicodeBold" if "ATSUnicodeBold" in pdfmetrics.getRegisteredFontNames() else "ATSUnicodeRegular"
        font_italic = "ATSUnicodeItalic" if "ATSUnicodeItalic" in pdfmetrics.getRegisteredFontNames() else "ATSUnicodeRegular"

    return font_regular, font_bold, font_italic


def save_improved_cv_text(improved_text: str, original_ext: str, ts: str) -> tuple[str, Path, str]:
    ext = original_ext.lower()

    if ext == ".txt":
        out_name = f"cv_improved_{ts}.txt"
        out_path = GENERATED_DIR / out_name
        out_path.write_text(improved_text, encoding="utf-8")
        return out_name, out_path, "txt"

    if ext == ".docx" or ext == ".doc":
        try:
            from docx import Document
            out_name = f"cv_improved_{ts}.docx"
            out_path = GENERATED_DIR / out_name
            doc = Document()
            for line in improved_text.splitlines():
                doc.add_paragraph(line)
            doc.save(str(out_path))
            return out_name, out_path, "docx"
        except Exception as e:
            raise RuntimeError(f"Nie udało się zapisać DOCX: {e}")

    if ext == ".pdf":
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

            out_name = f"cv_improved_{ts}.pdf"
            out_path = GENERATED_DIR / out_name
            font_regular, font_bold, _ = _build_unicode_fonts_for_pdf()

            doc = SimpleDocTemplate(
                str(out_path),
                pagesize=A4,
                leftMargin=1.8 * cm,
                rightMargin=1.8 * cm,
                topMargin=1.5 * cm,
                bottomMargin=1.5 * cm,
            )
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle("Title", parent=styles["Heading1"], fontName=font_bold, fontSize=14)
            body_style = ParagraphStyle("Body", parent=styles["BodyText"], fontName=font_regular, fontSize=10, leading=14)

            story = [Paragraph("CV — Zaktualizowana wersja", title_style), Spacer(1, 10)]
            for line in improved_text.splitlines():
                story.append(Paragraph(html.escape(line) if line else "&nbsp;", body_style))
            doc.build(story)
            return out_name, out_path, "pdf"
        except Exception as e:
            raise RuntimeError(f"Nie udało się zapisać PDF: {e}")

    raise RuntimeError(f"Nieobsługiwany format pliku: {original_ext}")


def save_improved_cv(
    original_path: Path,
    improved_text: str,
    analysis: dict,
    output_format: str,
    ts: str,
) -> tuple[str, Path, str, str]:
    ext = original_path.suffix.lower()
    requested = (output_format or "").strip().lower()
    if not requested:
        requested = ext.lstrip(".")

    if requested == "doc":
        requested = "docx"

    if requested == "canva":
        requested = "html"

    html_name = ""

    if ext == ".pdf":
        layout = extract_pdf_layout(original_path)
        for page in layout:
            for line in page["lines"]:
                line["text"] = improve_line_text(str(line["text"]), analysis)

        html_name, _ = build_html_from_layout(layout, analysis, ts)

        if requested == "html":
            out_path = HTML_PREVIEW_DIR / html_name
            return html_name, out_path, "html", html_name

        if requested == "pdf":
            out_name = f"cv_improved_{ts}.pdf"
            out_path = GENERATED_DIR / out_name
            render_pdf_from_layout(layout, out_path)
            return out_name, out_path, "pdf", html_name

        if requested == "docx":
            out_name = f"cv_improved_{ts}.docx"
            out_path = GENERATED_DIR / out_name
            from docx import Document

            doc = Document()
            for page in layout:
                for line in page["lines"]:
                    doc.add_paragraph(str(line["text"]))
                doc.add_page_break()
            doc.save(str(out_path))
            return out_name, out_path, "docx", html_name

        raise RuntimeError(f"Nieobsługiwany format wyjściowy: {requested}")

    if requested in {"txt", "docx", "pdf"}:
        out_name, out_path, out_fmt = save_improved_cv_text(improved_text, f".{requested}", ts)

        html_body = "\n".join(f"<p>{html.escape(line) or '&nbsp;'}</p>" for line in improved_text.splitlines())
        html_name = f"cv_improved_preview_{ts}.html"
        html_path = HTML_PREVIEW_DIR / html_name
        html_path.write_text(
            "<!doctype html><html lang=\"pl\"><head><meta charset=\"utf-8\"><title>CV preview</title></head><body>"
            + html_body
            + "</body></html>",
            encoding="utf-8",
        )
        return out_name, out_path, out_fmt, html_name

    if requested == "html":
        html_name = f"cv_improved_preview_{ts}.html"
        html_path = HTML_PREVIEW_DIR / html_name
        html_body = "\n".join(f"<p>{html.escape(line) or '&nbsp;'}</p>" for line in improved_text.splitlines())
        html_path.write_text(
            "<!doctype html><html lang=\"pl\"><head><meta charset=\"utf-8\"><title>CV preview</title></head><body>"
            + html_body
            + "</body></html>",
            encoding="utf-8",
        )
        return html_name, html_path, "html", html_name

    raise RuntimeError(f"Nieobsługiwany format wyjściowy: {requested}")


def inline_preview_assets(html_text: str, html_path: Path) -> str:
    return preview_inline_preview_assets(html_text, html_path)


def export_html_preview_file(source_html: Path, output_format: str, ts: str) -> tuple[str, Path, str]:
    return preview_export_html_preview_file(source_html, output_format, ts)


def convert_faithful_template(source_file: Path, output_format: str, ts: str) -> dict:
    result = preview_convert_faithful_template(source_file, output_format, ts)
    if result.get("success") is not False and result.get("html_preview"):
        _manifest_register(source_file.name, result["html_preview"], ts)
    return result


@app.route("/api/html-previews")
def list_html_previews():
    return jsonify({"success": True, "items": list_html_preview_items()})


@app.route("/api/html-preview-content/<path:filename>")
def html_preview_content(filename):
    try:
        safe, fp = read_html_preview_content(filename)
    except FileNotFoundError:
        return jsonify({"success": False, "error": "Podgląd HTML nie istnieje"}), 404
    try:
        content = fp.read_text(encoding="utf-8")
    except Exception as e:
        return jsonify({"success": False, "error": f"Nie udało się odczytać HTML: {e}"}), 500
    return jsonify({"success": True, "filename": safe, "content": content})


@app.route("/api/html-preview-content/<path:filename>", methods=["POST"])
def save_html_preview_content(filename):
    try:
        safe, fp = read_html_preview_content(filename)
    except FileNotFoundError:
        return jsonify({"success": False, "error": "Podgląd HTML nie istnieje"}), 404

    data = request.get_json(silent=True) or {}
    content = data.get("content", "")
    if not isinstance(content, str) or not content.strip():
        return jsonify({"success": False, "error": "Treść HTML nie może być pusta"}), 400

    if "<html" not in content.lower() and "<!doctype html" not in content.lower():
        return jsonify({"success": False, "error": "Treść nie wygląda na pełny dokument HTML"}), 400

    try:
        fp.write_text(content, encoding="utf-8")
    except Exception as e:
        return jsonify({"success": False, "error": f"Nie udało się zapisać HTML: {e}"}), 500

    return jsonify({"success": True, "filename": safe, "preview_url": f"/api/html-preview/{safe}"})


@app.route("/api/html-preview/<path:filename>", methods=["DELETE"])
def delete_html_preview(filename):
    safe = Path(filename).name
    if not safe or ".." in safe:
        return jsonify({"success": False, "error": "Nieprawidłowa nazwa pliku"}), 400

    deleted: list[str] = []

    # 1. Remove preview HTML
    preview_path = HTML_PREVIEW_DIR / safe
    if preview_path.exists():
        preview_path.unlink()
        deleted.append(str(preview_path.name))

    # 2. Find linked upload via manifest
    manifest = _manifest_load()
    upload_filename: str | None = None
    for upl, entry in manifest.items():
        if safe in entry.get("previews", []):
            upload_filename = upl
            break

    timestamps: list[str] = []
    if upload_filename:
        entry = manifest[upload_filename]
        # Collect timestamps for this upload (used to find assets)
        timestamps = entry.get("ts", [])
        # Remove from manifest
        entry["previews"] = [p for p in entry["previews"] if p != safe]
        if not entry["previews"]:
            del manifest[upload_filename]
        _manifest_save(manifest)

        # 3. Delete original upload if no more previews remain
        upl_path = UPLOAD_DIR / upload_filename
        if upl_path.exists():
            upl_path.unlink()
            deleted.append(upload_filename)

    # 4. Remove conversion assets dirs matching the safe filename or known timestamps
    stem = Path(safe).stem  # e.g. cv_faithful_20260610_143052
    ts_candidates: set[str] = set(timestamps)
    # Also extract ts from the preview filename itself (cv_faithful_YYYYMMDD_HHMMSS)
    m = re.search(r'_(\d{8}_\d{6})', stem)
    if m:
        ts_candidates.add(m.group(1))
    for ts_str in ts_candidates:
        asset_dir = CONVERSION_ASSETS_DIR / f"conv_{ts_str}"
        if asset_dir.exists():
            import shutil as _shutil
            _shutil.rmtree(str(asset_dir), ignore_errors=True)
            deleted.append(f"conv_{ts_str}/")
        for img_path in list(EXTRACTED_IMAGES_DIR.glob(f"cv_photo_{ts_str}_*")):
            img_path.unlink(missing_ok=True)
            deleted.append(img_path.name)

    return jsonify({"success": True, "deleted": deleted})


def _extract_ddg_results(html_text: str) -> list[dict]:
    results: list[dict] = []
    anchor_pattern = re.compile(
        r'<a[^>]*class="result__a"[^>]*href="([^"]+)"[^>]*>(.*?)</a>',
        re.IGNORECASE | re.DOTALL,
    )

    for href, title_html in anchor_pattern.findall(html_text):
        title = re.sub(r"<[^>]+>", "", title_html)
        title = html.unescape(title).strip()
        url = html.unescape(href).strip()
        if not url or not title:
            continue

        parsed = urllib.parse.urlparse(url)
        if parsed.netloc.endswith("duckduckgo.com") and parsed.path.startswith("/l/"):
            q = urllib.parse.parse_qs(parsed.query)
            if "uddg" in q and q["uddg"]:
                url = urllib.parse.unquote(q["uddg"][0])

        results.append({"title": title, "url": url})

    return results


def search_matching_jobs(cv_text: str, job_description: str = "") -> dict:
    cv_low = cv_text.lower()
    jd_low = (job_description or "").lower()

    role = "Senior PHP Developer Symfony"
    if "python" in jd_low and "php" not in jd_low:
        role = "Senior Python Developer"

    signal_terms = []
    for term in ["Symfony", "Redis", "RabbitMQ", "ElasticSearch", "MySQL", "Kubernetes", "Akamai"]:
        if term.lower() in jd_low or term.lower() in cv_low:
            signal_terms.append(term)

    query_terms = " ".join(signal_terms[:4])
    domain_filter = " OR ".join([f"site:{d}" for d in POPULAR_JOB_DOMAINS])
    query = f"{role} {query_terms} ({domain_filter})"

    ddg_url = "https://duckduckgo.com/html/?q=" + urllib.parse.quote_plus(query)
    req = urllib.request.Request(
        ddg_url,
        headers={
            "User-Agent": "Mozilla/5.0 (compatible; CV-ATS-Analyzer/1.0)",
        },
    )

    jobs = []
    try:
        with urllib.request.urlopen(req, timeout=12) as resp:
            content = resp.read().decode("utf-8", errors="ignore")
            raw_results = _extract_ddg_results(content)

        seen = set()
        for item in raw_results:
            url = item.get("url", "")
            if not url.startswith("http"):
                continue
            domain = urllib.parse.urlparse(url).netloc.lower()
            if domain.startswith("www."):
                domain = domain[4:]
            if not any(d in domain for d in POPULAR_JOB_DOMAINS):
                continue
            if url in seen:
                continue
            seen.add(url)
            jobs.append({
                "title": item.get("title", "Oferta pracy"),
                "url": url,
                "source": "DuckDuckGo / public search",
                "domain": domain,
                "snippet": "Sprawdź szczegóły oferty bezpośrednio na portalu.",
            })
            if len(jobs) >= 12:
                break
    except Exception:
        jobs = []

    if not jobs:
        encoded = urllib.parse.quote_plus(f"{role} {query_terms}".strip())
        jobs = [
            {
                "title": "Wyniki: Pracuj.pl",
                "url": f"https://www.pracuj.pl/praca/{encoded}",
                "source": "Portal jobs",
                "domain": "pracuj.pl",
                "snippet": "Brak bezpośrednich wyników z agregatora; otwórz wyniki wyszukiwania portalu.",
            },
            {
                "title": "Wyniki: JustJoin.it",
                "url": f"https://justjoin.it/job-offers/all-locations?keyword={encoded}",
                "source": "Portal jobs",
                "domain": "justjoin.it",
                "snippet": "Wyniki wyszukiwania ofert IT na JustJoin.it.",
            },
            {
                "title": "Wyniki: No Fluff Jobs",
                "url": f"https://nofluffjobs.com/pl/jobs?criteria=keyword%3D{encoded}",
                "source": "Portal jobs",
                "domain": "nofluffjobs.com",
                "snippet": "Wyniki wyszukiwania ofert IT na No Fluff Jobs.",
            },
            {
                "title": "Wyniki: LinkedIn Jobs",
                "url": f"https://www.linkedin.com/jobs/search/?keywords={encoded}",
                "source": "Portal jobs",
                "domain": "linkedin.com",
                "snippet": "Wyniki wyszukiwania ofert na LinkedIn Jobs.",
            },
        ]

    return {
        "query": query,
        "jobs": jobs,
    }


def parse_cv_for_template(cv_text: str) -> dict:
    lines = [l.strip() for l in cv_text.split("\n") if l.strip()]
    data: dict = {
        "name":    lines[0] if lines else "Jan Kowalski",
        "contact": [],
    }
    emails    = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', cv_text)
    phones    = re.findall(r'(?:\+\d{1,3}[\s\-]?)?\d[\d\s\-]{7,14}\d', cv_text)
    linkedin  = re.findall(r'linkedin\.com/in/[\w\-]+', cv_text, re.I)
    github    = re.findall(r'github\.com/[\w\-]+', cv_text, re.I)

    if emails:   data["contact"].append(f"✉ {emails[0]}")
    if phones:   data["contact"].append(f"☎ {phones[0].strip()}")
    if linkedin: data["contact"].append(f"in {linkedin[0]}")
    if github:   data["contact"].append(f"⌥ {github[0]}")
    return data


COLOR_SCHEMES = {
    "klasyczny":    {"primary": "#1a1a2e", "accent": "#16213e",  "line": "#c0c0c0"},
    "nowoczesny":   {"primary": "#0f4c75", "accent": "#1b262c",  "line": "#bbe1fa"},
    "techniczny":   {"primary": "#00b894", "accent": "#2d3436",  "line": "#dfe6e9"},
    "minimalny":    {"primary": "#2c2c2c", "accent": "#555555",  "line": "#e0e0e0"},
    "ats_friendly": {"primary": "#000000", "accent": "#333333",  "line": "#aaaaaa"},
}


def generate_pdf(parsed: dict, template_type: str, output_path: Path) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors as rl
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle,
        )
        from reportlab.lib.units import cm
        from reportlab.lib.enums import TA_LEFT
    except ImportError:
        raise RuntimeError("reportlab nie jest zainstalowany. Uruchom: pip install reportlab")

    scheme  = COLOR_SCHEMES.get(template_type, COLOR_SCHEMES["klasyczny"])
    def c(h: str):
        h = h.lstrip("#")
        return rl.Color(int(h[0:2], 16) / 255, int(h[2:4], 16) / 255, int(h[4:6], 16) / 255)

    primary = c(scheme["primary"])
    accent  = c(scheme["accent"])
    line_c  = c(scheme["line"])

    doc = SimpleDocTemplate(
        str(output_path), pagesize=A4,
        leftMargin=1.8 * cm, rightMargin=1.8 * cm,
        topMargin=1.5 * cm,  bottomMargin=1.5 * cm,
    )
    base = getSampleStyleSheet()

    # Register Unicode-capable fonts to support Polish characters in generated PDF.
    font_regular, font_bold, font_italic = _build_unicode_fonts_for_pdf()

    def S(n, **kw): return ParagraphStyle(n, parent=base["Normal"], **kw)

    name_s    = S("N",  fontSize=22, fontName=font_bold,   textColor=primary, spaceAfter=2)
    contact_s = S("C",  fontSize=9,  fontName=font_regular,textColor=accent,  spaceAfter=10)
    sec_s     = S("Sc", fontSize=10, fontName=font_bold,   textColor=primary, spaceBefore=14, spaceAfter=4)
    body_s    = S("B",  fontSize=10, fontName=font_regular,textColor=rl.HexColor("#333333"), spaceAfter=3, leading=15)
    sub_s     = S("Su", fontSize=9,  fontName=font_regular,textColor=accent,  spaceAfter=4)
    ital_s    = S("It", fontSize=10, fontName=font_italic, textColor=rl.HexColor("#666666"), spaceAfter=3, leading=15)
    job_s     = S("J",  fontSize=11, fontName=font_bold,   textColor=primary, spaceAfter=2)
    key_s     = S("K",  fontSize=9,  fontName=font_bold,   textColor=primary)
    val_s     = S("V",  fontSize=9,  fontName=font_regular,textColor=rl.HexColor("#444444"))

    def HR():
        return HRFlowable(width="100%", thickness=0.7, color=line_c, spaceAfter=6, spaceBefore=2)

    def section(title):
        return [Paragraph(title.upper(), sec_s), HR()]

    story = []

    # Header
    name    = parsed.get("name", "Imię i Nazwisko")
    contact = parsed.get("contact") or ["twoj@email.com", "+48 000 000 000", "linkedin.com/in/profil"]
    story.append(Paragraph(name, name_s))
    story.append(Paragraph("  |  ".join(contact), contact_s))
    story.append(HRFlowable(width="100%", thickness=2, color=primary, spaceAfter=10))

    # Summary
    story += section("Podsumowanie Zawodowe")
    story.append(Paragraph(
        "[Senior PHP Developer z X+ lat doświadczenia w budowie skalowalnych platform e-commerce — "
        "Symfony, ElasticSearch, RabbitMQ, Redis. Entuzjasta AI-driven development "
        "i podejścia FAST / iteracyjnego. Doświadczenie z high-traffic systemami (50k+ concurrent users).]",
        ital_s,
    ))

    # Experience
    story += section("Doświadczenie Zawodowe")
    exp = [
        {
            "title": "Senior PHP Developer",
            "company": "[Nazwa Firmy] — Gdańsk / Remote",
            "period": "2021 – Obecnie",
            "bullets": [
                "Architektura i implementacja mikroserwisów w Symfony 6.x — SOLID, DDD, CQRS",
                "Profilowanie wydajności z BlackFire Profiler — redukcja czasu odpowiedzi o 40%",
                "Asynchroniczne pipeline'y: RabbitMQ + Dead Letter Exchanges + consumer pools",
                "ElasticSearch — full-text search + faceted navigation dla 10M+ produktów",
                "AI-driven development: GitHub Copilot, ChatGPT (30% szybsze feature delivery)",
            ],
        },
        {
            "title": "PHP Developer",
            "company": "[Poprzednia Firma] — Warszawa",
            "period": "2018 – 2021",
            "bullets": [
                "Rozwój platform e-commerce na Symfony + Doctrine ORM, integracja z systemami płatności",
                "Redis cache strategy — session handling, rate limiting, pub/sub, Lua scripting",
                "PHPUnit + Behat — pokrycie testami do 85%, TDD w codziennej pracy",
                "Code review, mentoring 2 junior developerów, dokumentacja techniczna",
            ],
        },
    ]
    for e in exp:
        story.append(Paragraph(f"{e['title']}  ·  {e['period']}", job_s))
        story.append(Paragraph(e["company"], sub_s))
        for b in e["bullets"]:
            story.append(Paragraph(f"▸  {b}", body_s))
        story.append(Spacer(1, 8))

    # Skills
    story += section("Umiejętności Techniczne")
    skills = [
        ("Backend",       "PHP 8.2, Symfony 7, Doctrine ORM, PHPUnit, Behat, BlackFire"),
        ("Bazy danych",   "MySQL 8, MongoDB, ElasticSearch 8, Redis 7"),
        ("Message Queue", "RabbitMQ (DLX, consumer groups, prefetch QoS)"),
        ("DevOps",        "Docker, Kubernetes, Jenkins, Linux, Git, CI/CD"),
        ("CDN / Cache",   "Varnish, Akamai, Redis, OPcache"),
        ("AI Tools",      "GitHub Copilot, ChatGPT, Claude — codzienne użytkowanie"),
        ("Wzorce",        "SOLID, DDD, CQRS, Repository, Factory, Observer, Command Bus"),
    ]
    for k, v in skills:
        tbl = Table(
            [[Paragraph(k, key_s), Paragraph(v, val_s)]],
            colWidths=[3.8 * cm, None], hAlign="LEFT",
        )
        tbl.setStyle(TableStyle([
            ("VALIGN",       (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING",  (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (0, -1),  8),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
        ]))
        story.append(tbl)
    story.append(Spacer(1, 4))

    # Education
    story += section("Wykształcenie")
    story.append(Paragraph("[Kierunek] | [Uczelnia] | [Rok ukończenia]", body_s))

    # Languages
    story += section("Języki")
    story.append(Paragraph(
        "🇵🇱 Polski — Native   |   🇬🇧 Angielski — B2/C1 (technical docs, daily Slack, code review)",
        body_s,
    ))

    # Certifications
    story += section("Certyfikaty i Kursy")
    story.append(Paragraph("[Symfony Certification / AWS Certified Developer / etc.]", ital_s))

    doc.build(story)


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/status")
def health():
    return jsonify({"status": "ok", "service": "CV ATS Analyzer", "version": "1.0.0"})


@app.route("/api/upload", methods=["POST"])
def upload_cv():
    if "cv" not in request.files:
        return jsonify({"success": False, "error": "Brak pliku CV (field: cv)"}), 400
    file = request.files["cv"]
    if not file.filename:
        return jsonify({"success": False, "error": "Nie wybrano pliku"}), 400

    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({"success": False, "error": f"Nieobsługiwany format: {ext}. Dozwolone: PDF, DOCX, TXT"}), 400

    stem     = re.sub(r"[^a-zA-Z0-9_\-]", "_", Path(file.filename).stem)[:60]
    filename = f"{stem}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
    dest     = UPLOAD_DIR / filename
    file.save(str(dest))

    return jsonify({
        "success":       True,
        "filename":      filename,
        "original_name": file.filename,
        "size":          dest.stat().st_size,
        "ext":           ext,
        "uploaded_at":   datetime.datetime.now().isoformat(),
    })


@app.route("/api/cv/<path:filename>")
def serve_cv(filename):
    safe = Path(filename).name
    fp   = UPLOAD_DIR / safe
    if not fp.exists():
        return jsonify({"error": "Nie znaleziono pliku"}), 404
    return send_file(fp)


@app.route("/api/analyze", methods=["POST"])
def analyze():
    data     = request.get_json(silent=True) or {}
    filename = data.get("filename", "").strip()
    job_description = data.get("job_description", "")
    if not filename:
        return jsonify({"success": False, "error": "Brak nazwy pliku (filename)"}), 400

    fp = UPLOAD_DIR / Path(filename).name
    if not fp.exists():
        return jsonify({"success": False, "error": "Plik nie istnieje na serwerze"}), 404

    cv_text = extract_text(fp)
    if not cv_text.strip() or cv_text.startswith("[ERROR"):
        return jsonify({"success": False, "error": f"Nie udało się wyekstrahować tekstu: {cv_text[:200]}"}), 422

    result = run_ats_analysis(cv_text, job_description=job_description)
    annotations = build_improvement_annotations(cv_text)
    preview_text = cv_text[:12000]
    text_summary = build_text_summary(cv_text, result, annotations)
    return jsonify({
        "success": True,
        "analysis": result,
        "preview_text": preview_text,
        "improvement_annotations": annotations,
        "text_summary": text_summary,
    })


@app.route("/api/improve-cv", methods=["POST"])
def improve_cv():
    data = request.get_json(silent=True) or {}
    filename = data.get("filename", "").strip()
    job_description = data.get("job_description", "")
    output_format = data.get("output_format", "")

    if not filename:
        return jsonify({"success": False, "error": "Brak nazwy pliku (filename)"}), 400

    fp = UPLOAD_DIR / Path(filename).name
    if not fp.exists():
        return jsonify({"success": False, "error": "Plik nie istnieje na serwerze"}), 404

    cv_text = extract_text(fp)
    if not cv_text.strip() or cv_text.startswith("[ERROR"):
        return jsonify({"success": False, "error": "Nie udało się odczytać tekstu CV do aktualizacji."}), 422

    analysis = run_ats_analysis(cv_text, job_description=job_description)
    improved_text = improve_cv_text(cv_text, analysis, job_description=job_description)

    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    try:
        out_name, _, output_format_used, html_preview_name = save_improved_cv(
            fp,
            improved_text,
            analysis,
            output_format=output_format,
            ts=ts,
        )
    except RuntimeError as e:
        return jsonify({"success": False, "error": str(e)}), 500

    return jsonify({
        "success": True,
        "download_filename": out_name,
        "output_format": output_format_used,
        "html_preview": html_preview_name,
        "html_preview_url": f"/api/html-preview/{html_preview_name}" if html_preview_name else "",
        "improved_preview": improved_text[:12000],
    })


@app.route("/api/convert-template", methods=["POST"])
def convert_template():
    data = request.get_json(silent=True) or {}
    filename = data.get("filename", "").strip()
    output_format = data.get("output_format", "html")

    if not filename:
        return jsonify({"success": False, "error": "Brak nazwy pliku (filename)"}), 400

    source = UPLOAD_DIR / Path(filename).name
    if not source.exists():
        return jsonify({"success": False, "error": "Plik nie istnieje na serwerze"}), 404

    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    try:
        payload = convert_faithful_template(source, output_format, ts)
    except RuntimeError as e:
        return jsonify({"success": False, "error": str(e)}), 500

    return jsonify({"success": True, **payload})


@app.route("/api/html-preview/<path:filename>")
def html_preview(filename):
    try:
        _, fp = read_html_preview_content(filename)
    except FileNotFoundError:
        return jsonify({"success": False, "error": "Podgląd HTML nie istnieje"}), 404
    try:
        content = inline_preview_assets(fp.read_text(encoding="utf-8"), fp)
    except Exception as e:
        return jsonify({"success": False, "error": f"Nie udało się przygotować podglądu HTML: {e}"}), 500
    return app.response_class(content, mimetype="text/html; charset=utf-8")


@app.route("/api/html-preview-export/<path:filename>")
def html_preview_export(filename):
    try:
        _, fp = read_html_preview_content(filename)
    except FileNotFoundError:
        return jsonify({"success": False, "error": "Podgląd HTML nie istnieje"}), 404

    output_format = request.args.get("format", "html")
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    try:
        out_name, out_path, _ = export_html_preview_file(fp, output_format, ts)
    except RuntimeError as e:
        return jsonify({"success": False, "error": str(e)}), 400
    except Exception as e:
        return jsonify({"success": False, "error": f"Nie udało się wyeksportować preview: {e}"}), 500

    return send_file(out_path, as_attachment=True, download_name=out_name)


@app.route("/api/search-jobs", methods=["POST"])
def search_jobs():
    data = request.get_json(silent=True) or {}
    filename = data.get("filename", "").strip()
    job_description = data.get("job_description", "")

    if not filename:
        return jsonify({"success": False, "error": "Brak nazwy pliku (filename)"}), 400

    fp = UPLOAD_DIR / Path(filename).name
    if not fp.exists():
        return jsonify({"success": False, "error": "Plik nie istnieje na serwerze"}), 404

    cv_text = extract_text(fp)
    if not cv_text.strip() or cv_text.startswith("[ERROR"):
        return jsonify({"success": False, "error": "Nie udało się odczytać tekstu CV do wyszukiwania ofert."}), 422

    payload = search_matching_jobs(cv_text, job_description=job_description)
    return jsonify({
        "success": True,
        "query": payload["query"],
        "jobs": payload["jobs"],
    })


@app.route("/api/templates")
def get_templates():
    return jsonify({
        "templates": [{"id": k, "name": v} for k, v in TEMPLATE_CATALOG.items()]
    })


@app.route("/api/generate-template", methods=["POST"])
def generate_template():
    data          = request.get_json(silent=True) or {}
    filename      = data.get("filename", "").strip()
    template_type = data.get("template_type", "klasyczny")

    if not filename:
        return jsonify({"success": False, "error": "Brak nazwy pliku"}), 400
    if template_type not in TEMPLATE_CATALOG:
        return jsonify({"success": False, "error": f"Nieznany typ szablonu: {template_type}"}), 400

    fp = UPLOAD_DIR / Path(filename).name
    if not fp.exists():
        return jsonify({"success": False, "error": "Plik nie istnieje"}), 404

    cv_text  = extract_text(fp)
    parsed   = parse_cv_for_template(cv_text)
    ts       = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = f"cv_{template_type}_{ts}.pdf"
    out_path = GENERATED_DIR / out_name

    try:
        generate_pdf(parsed, template_type, out_path)
    except RuntimeError as e:
        return jsonify({"success": False, "error": str(e)}), 500
    except Exception as e:
        return jsonify({"success": False, "error": f"Błąd generowania PDF: {e}"}), 500

    return jsonify({
        "success":           True,
        "download_filename": out_name,
        "template_name":     TEMPLATE_CATALOG[template_type],
    })


@app.route("/api/download/<path:filename>")
def download(filename):
    safe = Path(filename).name
    fp   = GENERATED_DIR / safe
    if not fp.exists():
        return jsonify({"error": "Plik nie istnieje"}), 404
    return send_file(fp, as_attachment=True, download_name=safe)


@app.errorhandler(413)
def too_large(_):
    return jsonify({"success": False, "error": "Plik za duży (maksymalnie 16 MB)"}), 413


if __name__ == "__main__":
    app.run(debug=True, host="127.0.0.1", port=9200)
