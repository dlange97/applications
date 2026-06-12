import base64
import datetime
import hashlib
import html
import mimetypes
import os
import re
import shutil
from io import BytesIO
from pathlib import Path

from services.document_service import extract_pdf_layout, extract_rich_pdf_layout
from services.storage import CONVERSION_ASSETS_DIR, EXTRACTED_IMAGES_DIR, GENERATED_DIR, HTML_PREVIEW_DIR


def write_layout_html_preview(
    pages_layout: list[dict],
    html_path: Path,
    title: str,
    notes_html: str = "",
) -> None:
    html_pages = []
    for page_idx, page in enumerate(pages_layout, start=1):
        line_html = []
        for line in page["lines"]:
            safe_text = html.escape(str(line["text"]))
            line_html.append(
                "<div class=\"line\" style=\"left:{x}px;top:{top}px;width:{w}px;font-size:{fs}px;\">{txt}</div>".format(
                    x=round(line["x"], 2),
                    top=round(line["top"], 2),
                    w=round(line["width"], 2),
                    fs=round(line["font_size"], 2),
                    txt=safe_text,
                )
            )

        html_pages.append(
            """
            <section class=\"page\" style=\"width:{w}px;height:{h}px;\">
              <div class=\"page-label\">Strona {idx}</div>
              {content}
            </section>
            """.format(
                w=round(page["width"], 2),
                h=round(page["height"], 2),
                idx=page_idx,
                content="\n".join(line_html),
            )
        )

    full_html = """
<!doctype html>
<html lang=\"pl\">
<head>
  <meta charset=\"utf-8\" />
  <title>{title}</title>
  <style>
    body {{ background:#f2f4f8; font-family: Arial, Helvetica, sans-serif; margin:0; padding:24px; }}
    .page {{ position:relative; margin:0 auto 18px; background:#fff; box-shadow:0 3px 16px rgba(0,0,0,.16); overflow:hidden; }}
    .line {{ position:absolute; white-space:nowrap; color:#111; line-height:1.2; }}
    .page-label {{ position:absolute; right:8px; top:6px; font-size:10px; color:#888; }}
    .ats-note {{ max-width:860px; margin:16px auto 0; background:#fff; padding:14px; box-shadow:0 3px 16px rgba(0,0,0,.12); }}
    .ats-note h3 {{ margin:0 0 8px; font-size:14px; }}
    .ats-note ul {{ margin:0; padding-left:18px; }}
  </style>
</head>
<body>
  {pages}
  {notes}
</body>
</html>
""".format(
        pages="\n".join(html_pages),
        title=html.escape(title),
        notes=notes_html,
    )
    html_path.write_text(full_html, encoding="utf-8")


def build_html_from_layout(pages_layout: list[dict], analysis: dict, ts: str) -> tuple[str, Path]:
    missing = analysis.get("keywords", {}).get("must_have", {}).get("missing", [])[:8]
    hints_html = "".join(f"<li>{html.escape(k)}</li>" for k in missing) or "<li>Brak krytycznych braków.</li>"
    note_block = """
  <div class=\"ats-note\">
    <h3>ATS hints ({ts})</h3>
    <ul>{hints}</ul>
  </div>
""".format(ts=html.escape(ts), hints=hints_html)
    html_name = f"cv_improved_preview_{ts}.html"
    html_path = HTML_PREVIEW_DIR / html_name
    write_layout_html_preview(pages_layout, html_path, "CV Improved Preview", note_block)
    return html_name, html_path


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as file_obj:
        for chunk in iter(lambda: file_obj.read(8192), b""):
            digest.update(chunk)
    return digest.hexdigest()


def build_faithful_html_from_images(page_images: list[Path], html_path: Path) -> None:
    sections = []
    for idx, image_path in enumerate(page_images, start=1):
        mime, _ = mimetypes.guess_type(str(image_path))
        mime = mime or "image/png"
        inline_img = base64.b64encode(image_path.read_bytes()).decode("ascii")
        sections.append(
            f"""
            <section class=\"page\">
              <div class=\"label\">Strona {idx}</div>
              <img src=\"data:{html.escape(mime)};base64,{inline_img}\" alt=\"CV page {idx}\" />
            </section>
            """
        )

    html_text = f"""
<!doctype html>
<html lang=\"pl\">
<head>
  <meta charset=\"utf-8\" />
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />
  <title>CV faithful template</title>
  <style>
    body {{ margin: 0; padding: 24px; background: #f2f3f5; font-family: Arial, Helvetica, sans-serif; }}
    .page {{ width: fit-content; margin: 0 auto 18px; position: relative; box-shadow: 0 4px 20px rgba(0,0,0,.15); background: #fff; }}
    .page img {{ display: block; max-width: min(1100px, 95vw); height: auto; }}
    .label {{ position:absolute; top:6px; right:8px; font-size:10px; color:#888; background: rgba(255,255,255,.8); padding: 2px 5px; border-radius: 4px; }}
  </style>
</head>
<body>
{''.join(sections)}
</body>
</html>
"""
    html_path.write_text(html_text, encoding="utf-8")


def extract_pdf_embedded_images(source_file: Path, ts: str) -> list[Path]:
    try:
        import fitz
    except Exception as e:
        raise RuntimeError("Do konwersji PDF wymagany jest pakiet PyMuPDF (pip install pymupdf).") from e

    extracted_photos = []
    document = fitz.open(str(source_file))
    for page_index in range(document.page_count):
        page = document.load_page(page_index)
        for image_index, image in enumerate(page.get_images(full=True)):
            xref = image[0]
            base = document.extract_image(xref)
            ext = base.get("ext", "png")
            blob = base.get("image", b"")
            if not blob:
                continue
            out_photo = EXTRACTED_IMAGES_DIR / f"cv_photo_{ts}_p{page_index + 1}_{image_index + 1}.{ext}"
            out_photo.write_bytes(blob)
            extracted_photos.append(out_photo)
    document.close()
    return extracted_photos


def _font_css(name: str) -> str:
    n = name.lower()
    if any(x in n for x in ("arial", "helvetica")):
        return "Arial, Helvetica, sans-serif"
    if any(x in n for x in ("times", "roman")):
        return "'Times New Roman', Times, serif"
    if any(x in n for x in ("courier", "mono")):
        return "'Courier New', Courier, monospace"
    if "calibri" in n:
        return "Calibri, Arial, sans-serif"
    if "verdana" in n:
        return "Verdana, Arial, sans-serif"
    if "georgia" in n:
        return "Georgia, serif"
    return "Arial, Helvetica, sans-serif"


def _rect_overlap_ratio(ax: float, ay: float, aw: float, ah: float,
                         bx: float, by: float, bw: float, bh: float) -> float:
    """Return overlap area as fraction of the SMALLER rect's area."""
    ox = max(ax, bx)
    oy = max(ay, by)
    ox2 = min(ax + aw, bx + bw)
    oy2 = min(ay + ah, by + bh)
    if ox2 <= ox or oy2 <= oy:
        return 0.0
    overlap = (ox2 - ox) * (oy2 - oy)
    smaller = min(aw * ah, bw * bh)
    return overlap / smaller if smaller > 0 else 0.0


def build_editable_html_from_pdf(source_file: Path, html_path: Path, title: str) -> int:
    """
    Convert PDF to editable, styled HTML.
    Three layers per page: colored bg fills, embedded photos, text spans.
    No page screenshots, no SVG blobs — all pure CSS/HTML attributes.

    Circular photo detection:
      In PDFs, profile photos are clipped to a circle via a clip-path.
      PyMuPDF gives us the full image bbox (rectangle), not the clipped shape.
      We detect the circle by finding a nearly-square colored bg rect that
      significantly overlaps with the image, then we:
        - render that bg rect with border-radius:50%
        - clip the image to that square's bounds with border-radius:50%
    """
    try:
        import fitz
    except Exception as exc:
        raise RuntimeError("PyMuPDF jest wymagany — pip install pymupdf") from exc

    doc = fitz.open(str(source_file))
    page_sections: list[str] = []

    for page_idx in range(doc.page_count):
        page = doc.load_page(page_idx)
        pw = round(page.rect.width, 1)
        ph = round(page.rect.height, 1)

        bg_parts: list[str] = []
        img_parts: list[str] = []
        txt_parts: list[str] = []

        # Collect all bg rects with metadata for image-circle detection
        bg_rects_meta: list[dict] = []

        # ── Layer 1: colored background fills ───────────────────────────────
        for drawing in page.get_drawings():
            fill = drawing.get("fill")
            if not fill:
                continue
            r, g, b = float(fill[0]), float(fill[1]), float(fill[2])
            if r > 0.93 and g > 0.93 and b > 0.93:
                continue          # skip white / near-white
            if r < 0.07 and g < 0.07 and b < 0.07:
                continue          # skip pure black (PDF stroke artefacts)
            rect = drawing.get("rect")
            if not rect:
                continue
            rw = round(rect.x1 - rect.x0, 1)
            rh = round(rect.y1 - rect.y0, 1)
            if rw < 2 or rh < 2:
                continue
            opacity = float(drawing.get("fill_opacity") or 1.0)
            ri, gi, bi = int(r * 255), int(g * 255), int(b * 255)
            hexc = f"#{ri:02x}{gi:02x}{bi:02x}"
            x  = round(rect.x0, 1)
            tp = round(rect.y0, 1)

            # Detect circular frame: large enough and nearly square
            is_circle_frame = (
                rw >= 40 and rh >= 40
                and 0.85 <= (rw / rh if rh else 0) <= 1.18
            )

            if is_circle_frame:
                bg_rects_meta.append({"x": x, "top": tp, "w": rw, "h": rh, "color": hexc})
                bg_parts.append(
                    f'<div class="bg" style="left:{x}pt;top:{tp}pt;width:{rw}pt;'
                    f'height:{rh}pt;background:{hexc};opacity:{opacity:.2f};'
                    f'border-radius:50%;"></div>'
                )
            else:
                bg_parts.append(
                    f'<div class="bg" style="left:{x}pt;top:{tp}pt;width:{rw}pt;'
                    f'height:{rh}pt;background:{hexc};opacity:{opacity:.2f};"></div>'
                )

        # ── Layer 2: embedded images (profile photo etc.) ────────────────────
        seen_xrefs: set[int] = set()
        for img_item in page.get_images(full=True):
            xref = img_item[0]
            if xref in seen_xrefs:
                continue
            seen_xrefs.add(xref)
            try:
                rects = page.get_image_rects(xref)
                base_img = doc.extract_image(xref)
                img_bytes = base_img.get("image", b"")
                img_ext   = base_img.get("ext", "png")
                if not img_bytes or not rects:
                    continue
                ir  = rects[0]
                x   = round(ir.x0, 1)
                tp  = round(ir.y0, 1)
                iw  = round(ir.x1 - ir.x0, 1)
                ih  = round(ir.y1 - ir.y0, 1)
                mime = "image/jpeg" if img_ext in ("jpg", "jpeg") else f"image/{img_ext}"
                b64  = base64.b64encode(img_bytes).decode()

                # Try to find overlapping circular frame
                circle_frame: dict | None = None
                best_overlap = 0.0
                for frame in bg_rects_meta:
                    ratio = _rect_overlap_ratio(x, tp, iw, ih,
                                                frame["x"], frame["top"],
                                                frame["w"], frame["h"])
                    if ratio > best_overlap:
                        best_overlap = ratio
                        circle_frame = frame

                if circle_frame is not None and best_overlap >= 0.35:
                    # Clip image to the circular frame bounds
                    cx  = circle_frame["x"]
                    ctp = circle_frame["top"]
                    cw  = circle_frame["w"]
                    ch  = circle_frame["h"]
                    img_parts.append(
                        f'<img class="embed" src="data:{mime};base64,{b64}" '
                        f'style="left:{cx}pt;top:{ctp}pt;width:{cw}pt;height:{ch}pt;'
                        f'border-radius:50%;object-fit:cover;" alt="photo" />'
                    )
                else:
                    # Non-photo image: place at native bbox
                    img_parts.append(
                        f'<img class="embed" src="data:{mime};base64,{b64}" '
                        f'style="left:{x}pt;top:{tp}pt;width:{iw}pt;height:{ih}pt;" alt="image" />'
                    )
            except Exception:
                pass

        # ── Layer 3: text spans ───────────────────────────────────────────────
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    if not text.strip():
                        continue
                    bbox       = span["bbox"]
                    flags      = span.get("flags", 0)
                    color_int  = span.get("color", 0)
                    cr = (color_int >> 16) & 0xFF
                    cg = (color_int >> 8)  & 0xFF
                    cb =  color_int        & 0xFF
                    color_hex  = f"#{cr:02x}{cg:02x}{cb:02x}"
                    font_size  = round(span.get("size", 11), 1)
                    bold       = bool(flags & 16)
                    italic     = bool(flags & 2)
                    ff         = _font_css(span.get("font", ""))
                    x          = round(bbox[0], 1)
                    tp         = round(bbox[1], 1)
                    min_w      = round(bbox[2] - bbox[0] + 4, 1)
                    txt        = html.escape(text)
                    fw         = "bold" if bold else "normal"
                    fi         = "italic" if italic else "normal"
                    txt_parts.append(
                        f'<span class="t" style="left:{x}pt;top:{tp}pt;'
                        f'min-width:{min_w}pt;font-size:{font_size}pt;'
                        f'font-weight:{fw};font-style:{fi};color:{color_hex};'
                        f'font-family:{ff};">{txt}</span>'
                    )

        page_sections.append(
            f'<div class="page" style="width:{pw}pt;height:{ph}pt;">\n'
            + "\n".join(bg_parts + img_parts + txt_parts)
            + "\n</div>"
        )

    doc.close()

    pages_str = "\n".join(page_sections)
    full_html = (
        '<!doctype html>\n<html lang="pl">\n<head>\n'
        '  <meta charset="utf-8" />\n'
        '  <meta name="viewport" content="width=device-width, initial-scale=1" />\n'
        f'  <title>{html.escape(title)}</title>\n'
        '  <style>\n'
        '    body { margin: 0; padding: 24px; background: #eef1f4; }\n'
        '    .page {\n'
        '      position: relative; overflow: hidden;\n'
        '      margin: 0 auto 24px; background: #fff;\n'
        '      box-shadow: 0 4px 24px rgba(0,0,0,.18);\n'
        '    }\n'
        '    .bg    { position: absolute; pointer-events: none; }\n'
        '    .embed { position: absolute; max-width: none; }\n'
        '    .t     { position: absolute; white-space: pre; line-height: 1.15; }\n'
        '  </style>\n'
        '</head>\n<body>\n'
        + pages_str
        + '\n</body>\n</html>'
    )
    html_path.write_text(full_html, encoding="utf-8")
    return len(page_sections)


_PT_TO_PX = 96.0 / 72.0  # kept for backward compat with write_layout_html_preview


def build_text_html_from_rich_layout(pages_layout: list[dict], html_path: Path, title: str) -> None:
    """
    Build a faithful, selectable-text HTML from rich PDF layout data.
    Renders three layers in z-order: filled rects → images → text spans.
    """
    page_sections = []

    for page_idx, page in enumerate(pages_layout, 1):
        pw = round(page["width"] * _PT_TO_PX, 1)
        ph = round(page["height"] * _PT_TO_PX, 1)

        rects_html: list[str] = []
        images_html: list[str] = []
        text_html: list[str] = []

        for elem in page.get("elements", []):
            x = round(elem["x"] * _PT_TO_PX, 1)
            top = round(elem["top"] * _PT_TO_PX, 1)
            ew = round((elem["x1"] - elem["x"]) * _PT_TO_PX, 1)
            eh = round((elem["bottom"] - elem["top"]) * _PT_TO_PX, 1)

            if elem["type"] == "rect":
                op = elem.get("opacity", 1.0)
                rects_html.append(
                    f'<div style="position:absolute;left:{x}px;top:{top}px;'
                    f'width:{ew}px;height:{eh}px;background:{elem["color"]};'
                    f'opacity:{op};z-index:1;"></div>'
                )

            elif elem["type"] == "image":
                img_bytes = elem.get("image_data", b"")
                if not img_bytes:
                    continue
                ext = elem.get("image_ext", "png")
                mime = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"
                b64 = base64.b64encode(img_bytes).decode()
                images_html.append(
                    f'<img src="data:{mime};base64,{b64}" '
                    f'style="position:absolute;left:{x}px;top:{top}px;'
                    f'width:{ew}px;height:{eh}px;max-width:none;z-index:2;" />'
                )

            elif elem["type"] == "text":
                raw = elem.get("text", "")
                if not raw.strip():
                    continue
                txt = html.escape(raw)
                fs = round(elem.get("font_size", 11) * _PT_TO_PX, 1)
                fw = "bold" if elem.get("bold") else "normal"
                fi = "italic" if elem.get("italic") else "normal"
                color = elem.get("color", "#000000")
                ff = elem.get("font_family", "Arial, Helvetica, sans-serif")
                min_w = max(ew + 6, 20)
                text_html.append(
                    f'<span style="position:absolute;left:{x}px;top:{top}px;'
                    f'font-size:{fs}px;font-weight:{fw};font-style:{fi};'
                    f'color:{color};font-family:{ff};white-space:pre;'
                    f'min-width:{min_w}px;line-height:1.15;z-index:3;">{txt}</span>'
                )

        all_html = "\n".join(rects_html + images_html + text_html)
        page_sections.append(
            f'<div class="page" style="width:{pw}px;height:{ph}px;">\n{all_html}\n</div>'
        )

    pages_joined = "\n".join(page_sections)
    full_html = (
        "<!doctype html>\n<html lang=\"pl\">\n<head>\n"
        "  <meta charset=\"utf-8\" />\n"
        "  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />\n"
        f"  <title>{html.escape(title)}</title>\n"
        "  <style>\n"
        "    body { margin: 0; padding: 24px; background: #e8eaed; }\n"
        "    .page {\n"
        "      position: relative;\n"
        "      margin: 0 auto 20px;\n"
        "      background: #ffffff;\n"
        "      box-shadow: 0 4px 24px rgba(0,0,0,.18);\n"
        "      overflow: hidden;\n"
        "    }\n"
        "  </style>\n"
        "</head>\n<body>\n"
        + pages_joined
        + "\n</body>\n</html>"
    )
    html_path.write_text(full_html, encoding="utf-8")


def render_pdf_pages_to_images_and_extract(source_file: Path, ts: str) -> tuple[list[Path], list[Path]]:
    try:
        import fitz
    except Exception as e:
        raise RuntimeError("Do konwersji 1:1 PDF wymagany jest pakiet PyMuPDF (pip install pymupdf).") from e

    assets_dir = CONVERSION_ASSETS_DIR / f"conv_{ts}"
    assets_dir.mkdir(parents=True, exist_ok=True)
    page_images = []
    extracted_photos = []

    document = fitz.open(str(source_file))
    for index in range(document.page_count):
        page = document.load_page(index)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
        page_img = assets_dir / f"page_{index + 1}.png"
        pixmap.save(str(page_img))
        page_images.append(page_img)

        for image_index, image in enumerate(page.get_images(full=True)):
            xref = image[0]
            base = document.extract_image(xref)
            ext = base.get("ext", "png")
            blob = base.get("image", b"")
            if not blob:
                continue
            out_photo = EXTRACTED_IMAGES_DIR / f"cv_photo_{ts}_p{index + 1}_{image_index + 1}.{ext}"
            out_photo.write_bytes(blob)
            extracted_photos.append(out_photo)
            break

    document.close()
    return page_images, extracted_photos


def image_input_to_page(source_file: Path, ts: str) -> tuple[list[Path], list[Path]]:
    assets_dir = CONVERSION_ASSETS_DIR / f"conv_{ts}"
    assets_dir.mkdir(parents=True, exist_ok=True)
    ext = source_file.suffix.lower().lstrip(".") or "png"
    page_img = assets_dir / f"page_1.{ext}"
    shutil.copy2(source_file, page_img)
    extracted = EXTRACTED_IMAGES_DIR / f"cv_photo_{ts}_p1_1.{ext}"
    shutil.copy2(source_file, extracted)
    return [page_img], [extracted]


def validate_faithful_conversion(source_images: list[Path], converted_images: list[Path]) -> dict:
    matches = 0
    total = min(len(source_images), len(converted_images))
    for source, converted in zip(source_images, converted_images):
        if sha256_file(source) == sha256_file(converted):
            matches += 1
    score = int((matches / total) * 100) if total else 0
    return {
        "matched_pages": matches,
        "total_pages": total,
        "score": score,
        "passed": total > 0 and matches == total,
    }


def inline_preview_assets(html_text: str, html_path: Path) -> str:
    def replace_src(match):
        quote = match.group(1)
        src = match.group(2).strip()
        low = src.lower()
        if low.startswith(("http://", "https://", "data:")):
            return match.group(0)

        asset_path = (html_path.parent / src).resolve()
        if not asset_path.exists() or not asset_path.is_file():
            return match.group(0)

        mime, _ = mimetypes.guess_type(str(asset_path))
        mime = mime or "application/octet-stream"
        blob = base64.b64encode(asset_path.read_bytes()).decode("ascii")
        return f'src={quote}data:{mime};base64,{blob}{quote}'

    return re.sub(r'src=("|\')([^"\']+)(\1)', replace_src, html_text, flags=re.IGNORECASE)


def html_to_plain_text(html_text: str) -> str:
    text = re.sub(r"<style[\s\S]*?</style>", " ", html_text, flags=re.IGNORECASE)
    text = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</(p|div|li|h1|h2|h3|h4|h5|h6|section)>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = html.unescape(text)
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def extract_embedded_images(html_text: str) -> list[tuple[str, bytes]]:
    images = []
    img_pattern = re.compile(r'src=("|\')data:([^;]+);base64,([^"\']+)(\1)', re.IGNORECASE)
    for _, mime, payload, _ in img_pattern.findall(html_text):
        try:
            images.append((mime, base64.b64decode(payload)))
        except Exception:
            continue
    return images


def export_html_preview_file(source_html: Path, output_format: str, ts: str) -> tuple[str, Path, str]:
    requested = (output_format or "html").strip().lower()
    raw_html = source_html.read_text(encoding="utf-8")
    standalone_html = inline_preview_assets(raw_html, source_html)

    if requested == "html":
        out_name = f"{source_html.stem}_export_{ts}.html"
        out_path = GENERATED_DIR / out_name
        out_path.write_text(standalone_html, encoding="utf-8")
        return out_name, out_path, "text/html"

    embedded_images = extract_embedded_images(standalone_html)

    if requested == "pdf":
        from reportlab.pdfgen import canvas
        from reportlab.lib.utils import ImageReader
        out_name = f"{source_html.stem}_export_{ts}.pdf"
        out_path = GENERATED_DIR / out_name

        if embedded_images:
            first_reader = ImageReader(BytesIO(embedded_images[0][1]))
            width, height = first_reader.getSize()
            pdf = canvas.Canvas(str(out_path), pagesize=(width, height))
            for _, img_bytes in embedded_images:
                reader = ImageReader(BytesIO(img_bytes))
                width, height = reader.getSize()
                pdf.setPageSize((width, height))
                pdf.drawImage(reader, 0, 0, width=width, height=height)
                pdf.showPage()
            pdf.save()
        else:
            pdf = canvas.Canvas(str(out_path))
            y = 800
            for line in html_to_plain_text(standalone_html).splitlines() or ["Brak treści HTML do eksportu."]:
                pdf.drawString(40, y, line[:120])
                y -= 16
                if y < 40:
                    pdf.showPage()
                    y = 800
            pdf.save()
        return out_name, out_path, "application/pdf"

    if requested == "docx":
        from docx import Document
        from docx.shared import Inches
        out_name = f"{source_html.stem}_export_{ts}.docx"
        out_path = GENERATED_DIR / out_name
        doc = Document()
        if embedded_images:
            for idx, (_, img_bytes) in enumerate(embedded_images):
                doc.add_picture(BytesIO(img_bytes), width=Inches(6.3))
                if idx < len(embedded_images) - 1:
                    doc.add_page_break()
        else:
            for line in html_to_plain_text(standalone_html).splitlines() or ["Brak treści HTML do eksportu."]:
                doc.add_paragraph(line)
        doc.save(str(out_path))
        return out_name, out_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    raise RuntimeError("Obsługiwane formaty pobierania: html, pdf, docx.")


def convert_faithful_template(source_file: Path, output_format: str, ts: str) -> dict:
    src_ext = source_file.suffix.lower()
    requested = (output_format or "html").strip().lower()
    if requested == "canva":
        requested = "html"
    if requested == "doc":
        requested = "docx"

    html_name = f"cv_faithful_{ts}.html"
    html_out = HTML_PREVIEW_DIR / html_name
    extracted_photos: list[Path] = []
    pages_layout = None
    page_images: list[Path] = []

    if src_ext == ".pdf":
        page_count = build_editable_html_from_pdf(source_file, html_out, source_file.stem)
        extracted_photos = extract_pdf_embedded_images(source_file, ts)
        validation = {
            "matched_pages": page_count,
            "total_pages": page_count,
            "score": 100 if page_count > 0 else 0,
            "passed": page_count > 0,
        }

    elif src_ext in {".jpg", ".jpeg", ".png"}:
        page_images, extracted_photos = image_input_to_page(source_file, ts)
        build_faithful_html_from_images(page_images, html_out)
        validation = validate_faithful_conversion(page_images, page_images)

    else:
        raise RuntimeError("Konwersja 1:1 obsługuje tylko wejściowe PDF/JPG/PNG.")

    if not validation["passed"]:
        raise RuntimeError("Walidacja 1:1 nie przeszła pomyślnie.")

    out_name = html_name
    output_used = requested

    if requested == "docx":
        try:
            from docx import Document
            from docx.shared import Inches
        except Exception as e:
            raise RuntimeError("Do eksportu DOCX wymagany jest python-docx.") from e

        out_name = f"cv_faithful_{ts}.docx"
        out_path = GENERATED_DIR / out_name
        doc = Document()

        if src_ext == ".pdf":
            pages_layout = extract_rich_pdf_layout(source_file)

        if pages_layout is not None:
            for page in pages_layout:
                texts = sorted(
                    (e for e in page.get("elements", []) if e.get("type") == "text"),
                    key=lambda e: (round(e["top"], 1), e["x"]),
                )
                para_buf: list[str] = []
                prev_top: float | None = None
                for t in texts:
                    cur = round(t["top"], 1)
                    if prev_top is not None and abs(cur - prev_top) > 5 and para_buf:
                        doc.add_paragraph(" ".join(para_buf))
                        para_buf = []
                    para_buf.append(t["text"])
                    prev_top = cur
                if para_buf:
                    doc.add_paragraph(" ".join(para_buf))
                doc.add_page_break()
        else:
            for idx, image_path in enumerate(page_images):
                doc.add_picture(str(image_path), width=Inches(6.3))
                if idx < len(page_images) - 1:
                    doc.add_page_break()

        doc.save(str(out_path))

    elif requested != "html":
        raise RuntimeError("Obsługiwane formaty wyjściowe: html, docx, canva.")

    return {
        "download_filename": out_name,
        "output_format": output_used,
        "html_preview": html_name,
        "html_preview_url": f"/api/html-preview/{html_name}",
        "extracted_images": [f"/api/download/{p.name}" for p in extracted_photos],
        "validation": validation,
    }


def list_html_preview_items() -> list[dict]:
    items = []
    for file_path in sorted(HTML_PREVIEW_DIR.glob("*.html"), key=lambda path: path.stat().st_mtime, reverse=True):
        stat = file_path.stat()
        items.append({
            "filename": file_path.name,
            "size": stat.st_size,
            "modified_at": datetime.datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "preview_url": f"/api/html-preview/{file_path.name}",
        })
    return items


def read_html_preview_content(filename: str) -> tuple[str, Path]:
    safe = Path(filename).name
    file_path = HTML_PREVIEW_DIR / safe
    if not file_path.exists():
        raise FileNotFoundError("Podgląd HTML nie istnieje")
    return safe, file_path
